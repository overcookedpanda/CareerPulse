import io

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _is_heading(line: str) -> bool:
    return (
        (line == line.upper() and len(line) > 2 and line.replace(" ", "").replace("&", "").replace("/", "").isalpha())
        or (line.endswith(":") and len(line) < 60 and not line.startswith("-"))
    )


def _is_subheading(line: str) -> bool:
    return (
        "|" in line
        or (any(c.isdigit() for c in line) and ("-" in line or "\u2013" in line) and len(line) < 120)
    )


def generate_resume_docx(resume_text: str, name: str = "") -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    core = doc.core_properties
    first_line = ""
    for line in resume_text.split("\n"):
        if line.strip():
            first_line = line.strip()
            break
    doc_name = name or first_line
    core.title = f"Resume - {doc_name}"
    core.author = doc_name

    is_first_line = True

    for raw_line in resume_text.split("\n"):
        line = raw_line.strip()

        if is_first_line and line:
            is_first_line = False
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)
            p.space_after = Pt(4)
            continue

        if is_first_line:
            continue

        if not line:
            doc.add_paragraph()
            continue

        if _is_heading(line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            p.space_before = Pt(6)
            p.space_after = Pt(2)
            # Add a bottom border via paragraph formatting
            pf = p.paragraph_format
            pf.space_after = Pt(2)
            continue

        if _is_subheading(line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            p.space_before = Pt(2)
            p.space_after = Pt(1)
            continue

        p = doc.add_paragraph(line)
        p.space_before = Pt(0)
        p.space_after = Pt(1)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_cover_letter_docx(cover_letter: str, company: str = "",
                                position: str = "") -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = "Cover Letter"
    if position and company:
        title = f"Cover Letter - {position} at {company}"
    elif position:
        title = f"Cover Letter - {position}"

    core = doc.core_properties
    core.title = title

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.space_after = Pt(12)

    for raw_line in cover_letter.split("\n"):
        line = raw_line.strip()
        if not line:
            doc.add_paragraph()
            continue

        is_greeting = (
            line.startswith("Dear ") or line.startswith("Sincerely")
            or line.startswith("Best regards") or line.startswith("Regards")
        )

        p = doc.add_paragraph()
        run = p.add_run(line)
        if is_greeting:
            run.bold = True
        p.space_after = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
